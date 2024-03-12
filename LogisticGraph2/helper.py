import docx

# doc = docx.Document("my_word_file.docx")
mydoc = docx.Document()
mydoc.add_paragraph("This is first paragraph of a MS Word file.")
mydoc.save("my_written_file.docx")

