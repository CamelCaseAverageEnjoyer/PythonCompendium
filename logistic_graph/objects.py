"""Гори в аду мфти"""

# from tkinter import *
# from tkinter import ttk
import random
import numpy as np
import docx  # dox as 

class Graph:
    def __init__(self, name: str, equalgraph: bool = True):
        self.name = name
        self.n = -1
        self.linkers = []
        
        # Параметры направленного графа
        self.equalgraph = equalgraph
        self.len_line = 3 if equalgraph else 4

        # Tk параметры
        self.width = 1200
        self.height = 800
        self.space = 150

        # Чтение файла txt
        text = "" if equalgraph else "_unequalgraph"
        text_tmp = "Сделай по нормальному файл блин!"
        f = open(f"storage/{name + text}.TXT", "r", encoding="utf-8")
        flag = 0
        for line in f:
            if flag == 3:
                lst = line.split('-')
                if equalgraph:
                    self.linkers += [int(lst[0]) - 1, int(lst[1]) - 1, int(lst[2])]
                else:
                    self.linkers += [int(lst[0]) - 1, int(lst[1]) - 1, int(lst[2]), int(lst[3])]
                if int(lst[0]) > self.n or int(lst[0]) < 1 or int(lst[1]) > self.n or int(lst[1]) < 1:
                    raise ValueError(f"{text_tmp} Номера точек должны быть от 1 до N!")
                if int(lst[2]) <= 0 or (not equalgraph and int(lst[3]) <= 0):
                    raise ValueError(f"{text_tmp} Ты где отрицательные расстояния видел?")
            if flag == 2 and 'Палочки:' in line:
                flag = 3
            if flag == 1:
                self.n = int(line)
                flag = 2
            if flag == 0 and 'Точки:' in line:
                flag = 1
        f.close()

        # Проверки на вшивость
        if flag == 0:
            raise NameError(f"{text_tmp} Задание точек идёт на следующей строке после 'Точки:'")
        if flag == 2:
            raise NameError(f"{text_tmp} Задание палочек идёт на следующей строке после 'Палочки:'")

        # Параметры после прочтения файла
        self.table_min_distance = [[0 if i == j else -1 for i in range(self.n)] for j in range(self.n)]
        self.table_way = [[' ' for i in range(self.n)] for j in range(self.n)]
        self.linkers = [[self.linkers[self.len_line * i + j] for j in range(self.len_line)] for i in range(int(len(self.linkers) // self.len_line))]

    def calculate(self):
        # Инициализация
        f = open("storage/counter.TXT", 'r')
        counter = int([line for line in f][0])
        f.close()
        f = open("storage/counter.TXT", 'w')
        f.write(f"{counter + 1}")
        f.close()
        doc = docx.Document()

        for i in range(self.n):  # Каждая строка таблицы
            # builder.writeln(f"##########Начало с узла {i + 1}##########")
            doc.add_heading(f"Начало с узла: {i + 1}", 2)
            node_remained = [j for j in range(self.n)]
            distances = [0 if i == j else 1e10 for j in range(self.n)]
            way = ['0' for _ in range(self.n)]
            node_remained.remove(i)
            k = i
            accumulation = 0
            while len(node_remained) > 0:
                # builder.writeln(f"Узел {k + 1}, накопленная длина: {int(accumulation)}")
                # print(f"node_remained: {node_remained}")
                text = f"Узел {k + 1}, u{k + 1}={int(accumulation)}, w{k + 1}={way[k]} \n"
                self.table_way[i][k] = way[k] if way[k] != '0' else ' '

                # Проверка соседних узлов
                for j in range(self.n):
                    for link in self.linkers:
                        if j in node_remained:
                            if link[0] == k and link[1] == j or \
                                    link[0] == j and link[1] == k:
                                numb = 3 if (not self.equalgraph and link[0] == k) else 2
                                # print(f"numb={numb}")
                                text += f"u{j + 1}=min(u{j + 1}; u{k + 1}+c{k + 1}{j + 1})=min(" \
                                        f"{distances[j] if distances[j] < 1e10 else '∞'}; {accumulation}+{link[numb]}) = "
                                if way[k] == '0':
                                    way[j] = str(j + 1)
                                elif distances[j] > 1e9:
                                    way[j] = way[k]
                                elif way[j] == '-1':
                                    way[j] = str(j + 1) if distances[j] < link[numb] + accumulation else way[k]
                                else:
                                    way[j] = way[j] if distances[j] < link[numb] + accumulation else way[k]
                                distances[j] = min(distances[j], link[numb] + accumulation)
                                text += f"{distances[j]} ({way[j]})\n"
                doc.add_paragraph(text)
                distances_remain = [distances[j] for j in node_remained]
                j_min = int(np.argmin(distances_remain))
                accumulation = int(np.min(distances_remain))
                k = node_remained[j_min]  # Дальнейший узел
                # print(f"distances: {distances}, distances_remain:{distances_remain}, j_min:{j_min}, k:{k}, "
                #       f"accumulation:{accumulation}")
                self.table_min_distance[i][k] = int(accumulation)
                # self.table_way[i][k] = way[j_min]
                node_remained.remove(k)
            self.table_way[i][k] = way[k]
        # print(f"self.table_min_distance: {self.table_min_distance}")
        # print(f"way: {self.table_way}")

        # Создание таблицы, сохранение\
        doc.add_paragraph('Матрица кратчайших расстояний:')
        table = doc.add_table(0, self.n + 1)
        table.style = 'Table Grid'
        for i in range(self.n + 1):
            cells = table.add_row().cells
            for j in range(self.n + 1):
                if i > 0 and j > 0:
                    cells[j].text = f"{self.table_min_distance[i - 1][j - 1]}"
                if i == 0 and j > 0:
                    cells[j].text = f"{j}"
                if i > 0 and j == 0:
                    cells[j].text = f"{i}"
        doc.add_paragraph('\nМатрица указателей:')
        table = doc.add_table(0, self.n + 1)
        table.style = 'Table Grid'
        for i in range(self.n + 1):
            cells = table.add_row().cells
            for j in range(self.n + 1):
                if i > 0 and j > 0:
                    cells[j].text = f"{self.table_way[i - 1][j - 1]}"
                if i == 0 and j > 0:
                    cells[j].text = f"{j}"
                if i > 0 and j == 0:
                    cells[j].text = f"{i}"
        tmp = "" if self.equalgraph else "unequalgraph"
        doc.save(f"word_files/{self.name}_{tmp}_{counter:04}.docx")

    def plot(self):
        root = Tk()
        root.title("На что я трачу свою жизнь")
        icon = PhotoImage(file="storage/question.png")
        root.iconphoto(False, icon)

        root.geometry(f"{self.width}x{self.height}+300+100")
        root.minsize(self.width, self.height)
        root.maxsize(self.width, self.height)

        def draw(points):
            # Линии
            for link in self.linkers:
                canvas.create_line(points[link[0]][0], points[link[0]][1], points[link[1]][0], points[link[1]][1],
                                   fill='#CCCCFF')
                canvas.create_text(int((points[link[0]][0] + points[link[1]][0]) // 2),
                                   int((points[link[0]][1] + points[link[1]][1]) // 2),
                                   text=f"{link[2]}", fill="#800080", font=('Helvetica', '20', 'bold'))

            # Кружочки
            for i in range(self.n):
                s = 20
                canvas.create_oval(points[i][0] - s, points[i][1] - s, points[i][0] + s, points[i][1] + s,
                                   fill='#6495ED')
                canvas.create_text(points[i][0], points[i][1] + int(s // 6), text=f"{i + 1}", fill="#000080",
                                   font=('Helvetica', '20', 'bold'))

        def redraw(nothing=None):
            canvas.delete('all')
            points = [[random.randint(self.space, self.width-self.space),
                       random.randint(self.space, self.height-self.space)] for _ in range(self.n)]
            draw(points)

        def draw_manual(nothing=None):
            canvas.delete('all')
            h = 50*4
            d = 35*4
            points = [[self.space + 0 * d, self.space + 1 * h],
                      [self.space + 1 * d, self.space + 0 * h],
                      [self.space + 1 * d, self.space + 2 * h],
                      [self.space + 2 * d, self.space + 1 * h],
                      [self.space + 3 * d, self.space + 0 * h],
                      [self.space + 3 * d, self.space + 2 * h],
                      [self.space + 4 * d, self.space + 1 * h],
                      [self.space + 5 * d, self.space + 0 * h],
                      [self.space + 5 * d, self.space + 2 * h],
                      [self.space + 6 * d, self.space + 1 * h]]
            draw(points)

        label = Label(text=f"Граф задачи: {self.name}")
        label.pack()
        btn = ttk.Button(text="Перерисовать", command=redraw)
        btn.pack()
        canvas = Canvas(bg="white", width=self.width, height=self.height-50)
        canvas.pack(anchor=CENTER, expand=1)
        root.bind('<Return>', redraw)
        root.bind('<Button-2>', draw_manual)
        redraw()

        root.mainloop()
